#!/usr/bin/env python3
"""提取简历文件中的文本内容，输出纯文本到 stdout。

用法:
    python extract_resume.py <文件路径>

支持格式:
    .docx   需要 python-docx（pip install python-docx）
    .pdf    优先 pandoc，其次 pdftotext
    .txt/.md 直接按 UTF-8 读取
"""
import os
import subprocess
import sys


def extract_docx(path):
    from docx import Document
    doc = Document(path)
    parts = []

    # 段落
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    # 表格（简历常用表格排版）
    for t in doc.tables:
        for row in t.rows:
            cells = [c.text.strip() for c in row.cells]
            # 合并单元格会重复出现，去重相邻相同值
            dedup = []
            for c in cells:
                if not dedup or dedup[-1] != c:
                    dedup.append(c)
            line = ' | '.join(c for c in dedup if c)
            if line:
                parts.append(line)

    return '\n'.join(parts)


def extract_pdf(path):
    # 方案一：pdftotext（poppler）
    try:
        r = subprocess.run(['pdftotext', path, '-'],
                           capture_output=True, text=True, timeout=60)
        if r.returncode == 0 and r.stdout.strip():
            return r.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired, OSError):
        pass

    # 方案二：Python 库兜底（pymupdf 优先，pdfplumber 次之）
    for lib in ('pymupdf', 'pdfplumber'):
        try:
            if lib == 'pymupdf':
                import pymupdf
                doc = pymupdf.open(path)
                text = '\n'.join(page.get_text('text') for page in doc)
                doc.close()
            else:
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    text = '\n'.join((p.extract_text() or '') for p in pdf.pages)
            if text.strip():
                return text
        except Exception:
            continue

    raise RuntimeError(
        'PDF 提取失败：本机缺少 pdftotext（poppler）或 Python PDF 库。'
        '请让用户直接粘贴简历文本，或安装 poppler 后重试。'
    )


def main():
    if len(sys.argv) < 2:
        print('用法: python extract_resume.py <文件路径>', file=sys.stderr)
        sys.exit(1)

    path = sys.argv[1]
    if not os.path.exists(path):
        print(f'文件不存在: {path}', file=sys.stderr)
        sys.exit(1)

    ext = os.path.splitext(path)[1].lower()

    if ext == '.docx':
        text = extract_docx(path)
    elif ext == '.pdf':
        text = extract_pdf(path)
    else:
        with open(path, encoding='utf-8-sig', errors='replace') as f:
            text = f.read()

    print(text)


if __name__ == '__main__':
    main()
